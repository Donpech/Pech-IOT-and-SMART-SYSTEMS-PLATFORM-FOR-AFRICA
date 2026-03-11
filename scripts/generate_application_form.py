#!/usr/bin/env python3
"""
PECH Group Holdings Ltd — Candidate Application Form Generator

Generates professional Excel (.xlsx) and Word (.docx) versions of the
PECH Candidate Application Form with PECH branding (Sky Blue #00BFFF,
Orange/Amber #F5A623), formatted tables, and print-ready layout.

Requirements:
    pip install openpyxl python-docx

Usage:
    python generate_application_form.py

Output:
    - PECH_CANDIDATE_APPLICATION_FORM.xlsx (in current directory)
    - PECH_CANDIDATE_APPLICATION_FORM.docx (in current directory)
"""

import os
import sys

# ---------------------------------------------------------------------------
# Check dependencies
# ---------------------------------------------------------------------------
MISSING = []
try:
    import openpyxl
    from openpyxl.styles import (
        Font, PatternFill, Alignment, Border, Side, Protection
    )
    from openpyxl.utils import get_column_letter
    from openpyxl.worksheet.datavalidation import DataValidation
except ImportError:
    MISSING.append("openpyxl")

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    import copy
except ImportError:
    MISSING.append("python-docx")

if MISSING:
    print(f"Missing dependencies: {', '.join(MISSING)}")
    print(f"Install with: pip install {' '.join(MISSING)}")
    sys.exit(1)

# ---------------------------------------------------------------------------
# Brand constants
# ---------------------------------------------------------------------------
PECH_BLUE = "00BFFF"
PECH_ORANGE = "F5A623"
PECH_DARK_BLUE = "0099CC"
PECH_DARK_BG = "1B2838"
WHITE = "FFFFFF"
FOOTER_TAGLINE = "PECH Group Holdings Ltd — Technology & Infrastructure Enablers for People"
LIGHT_GRAY = "F5F5F5"
BORDER_GRAY = "CCCCCC"

OUTPUT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# ============================================================================
# EXCEL GENERATION
# ============================================================================

def create_excel():
    """Generate the PECH Candidate Application Form as an Excel workbook."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Application Form"

    # Styles
    header_font = Font(name="Calibri", size=16, bold=True, color=WHITE)
    header_fill = PatternFill(start_color=PECH_DARK_BG, end_color=PECH_DARK_BG, fill_type="solid")
    section_font = Font(name="Calibri", size=12, bold=True, color=WHITE)
    section_fill = PatternFill(start_color=PECH_BLUE, end_color=PECH_BLUE, fill_type="solid")
    label_font = Font(name="Calibri", size=10, bold=True)
    value_font = Font(name="Calibri", size=10)
    note_font = Font(name="Calibri", size=9, italic=True, color="666666")
    thin_border = Border(
        left=Side(style="thin", color=BORDER_GRAY),
        right=Side(style="thin", color=BORDER_GRAY),
        top=Side(style="thin", color=BORDER_GRAY),
        bottom=Side(style="thin", color=BORDER_GRAY),
    )
    light_fill = PatternFill(start_color=LIGHT_GRAY, end_color=LIGHT_GRAY, fill_type="solid")
    wrap_alignment = Alignment(wrap_text=True, vertical="top")

    # Column widths
    ws.column_dimensions["A"].width = 5
    ws.column_dimensions["B"].width = 30
    ws.column_dimensions["C"].width = 40
    ws.column_dimensions["D"].width = 30
    ws.column_dimensions["E"].width = 25

    # Print setup
    ws.page_setup.paperSize = ws.PAPERSIZE_A4
    ws.page_setup.orientation = "portrait"
    ws.print_options.horizontalCentered = True
    ws.page_margins.left = 0.5
    ws.page_margins.right = 0.5
    ws.page_margins.top = 0.5
    ws.page_margins.bottom = 0.5

    row = 1

    # --- HEADER ---
    ws.merge_cells(f"A{row}:E{row}")
    cell = ws[f"A{row}"]
    cell.value = "PECH GROUP HOLDINGS LTD"
    cell.font = header_font
    cell.fill = header_fill
    cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[row].height = 35
    row += 1

    ws.merge_cells(f"A{row}:E{row}")
    cell = ws[f"A{row}"]
    cell.value = "CANDIDATE APPLICATION FORM — CONFIDENTIAL"
    cell.font = Font(name="Calibri", size=12, bold=True, color=PECH_ORANGE)
    cell.fill = header_fill
    cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[row].height = 25
    row += 1

    ws.merge_cells(f"A{row}:E{row}")
    cell = ws[f"A{row}"]
    cell.value = "Lagos, Nigeria | pechgroupholdings.tech | Ref: PECH-HR-CAF-2026-____"
    cell.font = Font(name="Calibri", size=9, color=WHITE)
    cell.fill = header_fill
    cell.alignment = Alignment(horizontal="center")
    ws.row_dimensions[row].height = 20
    row += 1

    # Helper functions
    def add_section_header(title):
        nonlocal row
        row += 1  # blank row
        ws.merge_cells(f"A{row}:E{row}")
        cell = ws[f"A{row}"]
        cell.value = title
        cell.font = section_font
        cell.fill = section_fill
        cell.alignment = Alignment(horizontal="left", vertical="center")
        ws.row_dimensions[row].height = 28
        row += 1

    def add_field(label, col_span=1, dropdown_values=None):
        nonlocal row
        ws[f"B{row}"].value = label
        ws[f"B{row}"].font = label_font
        ws[f"B{row}"].fill = light_fill
        ws[f"B{row}"].border = thin_border
        ws[f"B{row}"].alignment = wrap_alignment

        value_cell = f"C{row}"
        if col_span > 1:
            ws.merge_cells(f"C{row}:E{row}")
        ws[value_cell].border = thin_border
        ws[value_cell].font = value_font
        ws[value_cell].alignment = wrap_alignment
        ws[value_cell].protection = Protection(locked=False)

        if dropdown_values:
            dv = DataValidation(type="list", formula1=f'"{",".join(dropdown_values)}"', allow_blank=True)
            dv.error = "Please select from the dropdown list"
            dv.errorTitle = "Invalid Selection"
            ws.add_data_validation(dv)
            dv.add(ws[value_cell])

        row += 1

    def add_note(text):
        nonlocal row
        ws.merge_cells(f"B{row}:E{row}")
        ws[f"B{row}"].value = text
        ws[f"B{row}"].font = note_font
        ws[f"B{row}"].alignment = wrap_alignment
        row += 1

    def add_table_header(headers):
        nonlocal row
        for i, h in enumerate(headers):
            col = get_column_letter(i + 2)  # start from column B
            ws[f"{col}{row}"].value = h
            ws[f"{col}{row}"].font = Font(name="Calibri", size=9, bold=True, color=WHITE)
            ws[f"{col}{row}"].fill = PatternFill(start_color=PECH_DARK_BLUE, end_color=PECH_DARK_BLUE, fill_type="solid")
            ws[f"{col}{row}"].border = thin_border
            ws[f"{col}{row}"].alignment = Alignment(horizontal="center", wrap_text=True)
        row += 1

    def add_table_row(num_cols, num_rows=1):
        nonlocal row
        for _ in range(num_rows):
            for i in range(num_cols):
                col = get_column_letter(i + 2)
                ws[f"{col}{row}"].border = thin_border
                ws[f"{col}{row}"].font = value_font
                ws[f"{col}{row}"].alignment = wrap_alignment
                ws[f"{col}{row}"].protection = Protection(locked=False)
            row += 1

    # --- SECTION 1: PERSONAL INFORMATION ---
    add_section_header("SECTION 1: PERSONAL INFORMATION")
    add_field("Surname / Last Name", col_span=3)
    add_field("First Name", col_span=3)
    add_field("Middle Name", col_span=3)
    add_field("Date of Birth (DD/MM/YYYY)", col_span=3)
    add_field("Gender", dropdown_values=["Male", "Female", "Prefer not to say"])
    add_field("Marital Status", dropdown_values=["Single", "Married", "Divorced", "Widowed"])
    add_field("Nationality", col_span=3)
    add_field("State of Origin", col_span=3)
    add_field("Local Government Area (LGA)", col_span=3)
    add_field("Phone Number (Primary)", col_span=3)
    add_field("Phone Number (Secondary)", col_span=3)
    add_field("WhatsApp Number", col_span=3)
    add_field("Email Address", col_span=3)
    add_field("Residential Address", col_span=3)
    add_field("City", col_span=3)
    add_field("State", col_span=3)
    add_field("NIN (National ID Number)", col_span=3)

    # --- SECTION 2: POSITION APPLIED FOR ---
    add_section_header("SECTION 2: POSITION APPLIED FOR")
    add_field("Position / Job Title", col_span=3)
    add_field("Role ID (if known)", col_span=3)
    add_field("Department", dropdown_values=["Engineering", "Design", "Product & Business", "Operations", "Internship"])
    add_field("Employment Type", dropdown_values=["Full-Time", "Contract", "Internship"])
    add_field("Expected Monthly Salary (NGN)", col_span=3)
    add_field("Available Start Date", col_span=3)
    add_field("Notice Period (if employed)", col_span=3)
    add_field("How did you hear about us?", dropdown_values=["PECH Website", "LinkedIn", "Job Board", "Employee Referral", "University", "Social Media", "Other"])

    # --- SECTION 3: EDUCATIONAL BACKGROUND ---
    add_section_header("SECTION 3: EDUCATIONAL BACKGROUND")
    add_note("List all qualifications starting from the most recent")
    add_table_header(["Institution", "Qualification", "Course of Study", "Year"])
    add_table_row(4, num_rows=4)

    row += 1
    add_note("Professional Certifications")
    add_table_header(["Certification Name", "Issuing Body", "Year Obtained", "Certificate No."])
    add_table_row(4, num_rows=3)

    # --- SECTION 4: WORK EXPERIENCE ---
    add_section_header("SECTION 4: WORK EXPERIENCE")
    add_note("List your last 3 positions (most recent first). Write N/A if no experience.")

    for pos_num in range(1, 4):
        row += 1
        ws.merge_cells(f"B{row}:E{row}")
        ws[f"B{row}"].value = f"Position {pos_num}" + (" (Most Recent)" if pos_num == 1 else "")
        ws[f"B{row}"].font = Font(name="Calibri", size=10, bold=True, color=PECH_DARK_BLUE)
        row += 1
        add_field("Company / Organisation", col_span=3)
        add_field("Your Job Title", col_span=3)
        add_field("Employment Type", dropdown_values=["Full-Time", "Part-Time", "Contract", "Internship"])
        add_field("Start Date", col_span=3)
        add_field("End Date (or 'Present')", col_span=3)
        add_field("Monthly Salary (NGN)", col_span=3)
        add_field("Supervisor Name & Phone", col_span=3)
        add_field("Reason for Leaving", col_span=3)
        add_field("Key Responsibilities", col_span=3)

    # --- SECTION 5: TECHNICAL SKILLS ---
    add_section_header("SECTION 5: TECHNICAL SKILLS & COMPETENCIES")
    add_note("For technical roles. Non-technical applicants may skip to General Competencies below.")
    add_table_header(["Skill / Technology", "Proficiency Level", "Years of Experience", ""])
    for _ in range(8):
        ws[f"C{row}"].value = ""
        # Add proficiency dropdown
        dv = DataValidation(type="list", formula1='"Beginner,Intermediate,Advanced,Expert"', allow_blank=True)
        ws.add_data_validation(dv)
        dv.add(ws[f"C{row}"])
        for i in range(4):
            col = get_column_letter(i + 2)
            ws[f"{col}{row}"].border = thin_border
            ws[f"{col}{row}"].font = value_font
            ws[f"{col}{row}"].protection = Protection(locked=False)
        row += 1

    row += 1
    add_field("Portfolio Website URL", col_span=3)
    add_field("GitHub Profile URL", col_span=3)
    add_field("LinkedIn Profile URL", col_span=3)

    # --- SECTION 6: REFERENCES ---
    add_section_header("SECTION 6: REFERENCES (Minimum 2)")
    for ref_num in range(1, 3):
        ws.merge_cells(f"B{row}:E{row}")
        ws[f"B{row}"].value = f"Reference {ref_num}"
        ws[f"B{row}"].font = Font(name="Calibri", size=10, bold=True, color=PECH_DARK_BLUE)
        row += 1
        add_field("Full Name", col_span=3)
        add_field("Relationship", col_span=3)
        add_field("Organisation & Position", col_span=3)
        add_field("Phone Number", col_span=3)
        add_field("Email Address", col_span=3)
        row += 1

    # --- SECTION 7: GUARANTOR INFORMATION ---
    add_section_header("SECTION 7: GUARANTOR INFORMATION")
    add_note("Two guarantors required. Full Guarantor Forms will be completed upon offer.")
    for g_num in range(1, 3):
        ws.merge_cells(f"B{row}:E{row}")
        ws[f"B{row}"].value = f"Guarantor {g_num}"
        ws[f"B{row}"].font = Font(name="Calibri", size=10, bold=True, color=PECH_DARK_BLUE)
        row += 1
        add_field("Full Name", col_span=3)
        add_field("Occupation", col_span=3)
        add_field("Organisation / Employer", col_span=3)
        add_field("Phone Number", col_span=3)
        add_field("Residential Address", col_span=3)
        row += 1

    # --- SECTION 8: HEALTH & EMERGENCY ---
    add_section_header("SECTION 8: HEALTH & EMERGENCY INFORMATION")
    add_field("Blood Group", dropdown_values=["A+", "A-", "B+", "B-", "AB+", "AB-", "O+", "O-"])
    add_field("Genotype", dropdown_values=["AA", "AS", "SS", "AC", "SC"])
    add_field("Medical Conditions (if any)", col_span=3)
    add_field("Disability / Accommodation Needed", col_span=3)
    row += 1
    ws.merge_cells(f"B{row}:E{row}")
    ws[f"B{row}"].value = "Emergency Contact"
    ws[f"B{row}"].font = Font(name="Calibri", size=10, bold=True, color=PECH_DARK_BLUE)
    row += 1
    add_field("Full Name", col_span=3)
    add_field("Relationship", col_span=3)
    add_field("Phone Number", col_span=3)

    # --- SECTION 9: DECLARATIONS ---
    add_section_header("SECTION 9: ADDITIONAL DECLARATIONS")
    declarations = [
        ("Convicted of a criminal offence?", ["Yes", "No"]),
        ("Previously dismissed/terminated?", ["Yes", "No"]),
        ("Ongoing legal proceedings?", ["Yes", "No"]),
        ("Relatives at PECH Group?", ["Yes", "No"]),
        ("Currently employed elsewhere?", ["Yes", "No"]),
        ("Willing to relocate in Nigeria?", ["Yes", "No", "Negotiable"]),
        ("Willing to travel for work?", ["Yes", "No", "Domestic Only"]),
        ("Valid driver's license?", ["Yes", "No"]),
        ("Valid international passport?", ["Yes", "No"]),
        ("Authorized to work in Nigeria?", ["Yes", "No"]),
    ]
    for label, options in declarations:
        add_field(label, dropdown_values=options)

    add_field("If 'Yes' to any above, provide details:", col_span=3)

    # --- SECTION 10: CONSENT & SIGNATURE ---
    add_section_header("SECTION 10: DECLARATION & CONSENT")
    row += 1
    ws.merge_cells(f"B{row}:E{row}")
    ws[f"B{row}"].value = (
        "I declare that all information in this form is true and accurate. "
        "I consent to PECH Group Holdings Ltd processing my data per NDPA 2023. "
        "I understand false information may lead to disqualification or termination."
    )
    ws[f"B{row}"].font = Font(name="Calibri", size=9)
    ws[f"B{row}"].alignment = wrap_alignment
    ws.row_dimensions[row].height = 45
    row += 2

    add_field("Applicant's Full Name (Print)", col_span=3)
    add_field("Signature", col_span=3)
    add_field("Date (DD/MM/YYYY)", col_span=3)

    # --- SECTION 11: OFFICE USE ONLY ---
    add_section_header("SECTION 11: FOR OFFICE USE ONLY")
    add_note("To be completed by HR Department only")
    add_field("Application Reference No.", col_span=3)
    add_field("Date Received", col_span=3)
    add_field("Received By", col_span=3)
    add_field("Application Complete?", dropdown_values=["Yes", "No - Missing Documents"])
    add_field("Decision", dropdown_values=["Shortlist for Interview", "Waitlist", "Reject"])
    add_field("Reason for Decision", col_span=3)
    add_field("Approved By", col_span=3)

    # --- FOOTER ---
    row += 2
    ws.merge_cells(f"A{row}:E{row}")
    cell = ws[f"A{row}"]
    cell.value = FOOTER_TAGLINE
    cell.font = Font(name="Calibri", size=8, italic=True, color="999999")
    cell.alignment = Alignment(horizontal="center")

    # Save
    output_path = os.path.join(OUTPUT_DIR, "PECH_CANDIDATE_APPLICATION_FORM.xlsx")
    wb.save(output_path)
    print(f"Excel file generated: {output_path}")
    return output_path


# ============================================================================
# WORD GENERATION
# ============================================================================

def create_word():
    """Generate the PECH Candidate Application Form as a Word document.

    Features:
    - Colorful gradient section dividers (blue-to-orange)
    - Clickable checkbox form fields (☐ / ☑) that users can tick in Word
    - Professional PECH branding throughout
    """
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # Styles
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # --- Gradient divider colors (blue-to-orange, 4 segments) ---
    GRADIENT_COLORS = ["00BFFF", "0099CC", "F5A623", "E08A00"]

    def add_gradient_divider():
        """Add a colorful blue-to-orange gradient divider bar using a 1-row table."""
        table = doc.add_table(rows=1, cols=len(GRADIENT_COLORS))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Remove all borders and set cell shading
        for j, color in enumerate(GRADIENT_COLORS):
            cell = table.cell(0, j)
            # Set cell shading
            shading = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
            )
            cell._tc.get_or_add_tcPr().append(shading)
            # Set minimal height
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            pf = p.paragraph_format
            pf.line_spacing = Pt(4)
            # Remove cell borders
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '</w:tcBorders>'
            )
            tc_pr.append(borders)
        # Set row height
        tr = table.rows[0]._tr
        trPr = tr.get_or_add_trPr()
        trHeight = parse_xml(
            f'<w:trHeight {nsdecls("w")} w:val="80" w:hRule="exact"/>'
        )
        trPr.append(trHeight)
        return table

    def add_checkbox_sdt(paragraph, label, checked=False):
        """Add a clickable checkbox (Structured Document Tag) to a paragraph.

        This creates an actual Word checkbox content control that users can
        click to check/uncheck in Word (Yes ☑ / No ☐).
        """
        # Create the SDT (Structured Document Tag) for checkbox
        checked_val = "1" if checked else "0"
        symbol = "☑" if checked else "☐"

        sdt_xml = (
            f'<w:sdt {nsdecls("w")} {nsdecls("w14")}>'
            '  <w:sdtPr>'
            f'    <w14:checkbox><w14:checked w14:val="{checked_val}"/>'
            '      <w14:checkedState w14:val="2612" w14:font="MS Gothic"/>'
            '      <w14:uncheckedState w14:val="2610" w14:font="MS Gothic"/>'
            '    </w14:checkbox>'
            '  </w:sdtPr>'
            '  <w:sdtContent>'
            '    <w:r>'
            '      <w:rPr><w:rFonts w:ascii="MS Gothic" w:eastAsia="MS Gothic" w:hAnsi="MS Gothic"/>'
            '        <w:sz w:val="18"/><w:szCs w:val="18"/>'
            '      </w:rPr>'
            f'      <w:t>{symbol}</w:t>'
            '    </w:r>'
            '  </w:sdtContent>'
            '</w:sdt>'
        )
        sdt = parse_xml(sdt_xml)
        paragraph._p.append(sdt)
        # Add the label text after the checkbox
        run = paragraph.add_run(f" {label}   ")
        run.font.size = Pt(9)

    def add_checkbox_row(cell, options):
        """Fill a table cell with clickable checkboxes for each option."""
        p = cell.paragraphs[0]
        for option in options:
            add_checkbox_sdt(p, option)

    def add_header():
        """Add PECH branded header."""
        # Top gradient bar
        add_gradient_divider()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("PECH GROUP HOLDINGS LTD")
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0xBF, 0xFF)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("CANDIDATE APPLICATION FORM")
        run.font.size = Pt(14)
        run.font.bold = True

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("CONFIDENTIAL")
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xF5, 0xA6, 0x23)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Lagos, Nigeria | pechgroupholdings.tech")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        add_gradient_divider()

        p = doc.add_paragraph()
        run = p.add_run("Document Reference: ")
        run.font.bold = True
        run.font.size = Pt(9)
        p.add_run("PECH-HR-CAF-2026-________").font.size = Pt(9)
        p2 = doc.add_paragraph()
        run2 = p2.add_run("Form Version: 1.0  |  Date Issued: 10th March, 2026  |  Effective Year: 2026")
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    def add_section_heading(title):
        """Add a branded section heading with colorful gradient divider."""
        doc.add_paragraph()  # spacing
        add_gradient_divider()
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0xBF, 0xFF)

    def add_form_table(fields):
        """Add a form table with label-value pairs.

        If a value contains checkbox options (list of strings), actual
        clickable Word checkboxes are inserted instead of plain text.
        """
        table = doc.add_table(rows=len(fields), cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, field_data in enumerate(fields):
            label = field_data[0]
            default = field_data[1]
            # Label cell
            cell = table.cell(i, 0)
            cell.width = Cm(6)
            p = cell.paragraphs[0]
            run = p.add_run(label)
            run.font.bold = True
            run.font.size = Pt(9)
            # Value cell
            cell = table.cell(i, 1)
            cell.width = Cm(11)
            # Check if this is a checkbox field (list of options)
            if isinstance(default, list):
                add_checkbox_row(cell, default)
            else:
                p = cell.paragraphs[0]
                run = p.add_run(default)
                run.font.size = Pt(9)
        return table

    def add_note(text):
        """Add an italic note."""
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    def add_multi_row_table(headers, num_rows):
        """Add a table with headers and empty rows."""
        table = doc.add_table(rows=num_rows + 1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Header row with PECH dark blue background
        for j, h in enumerate(headers):
            cell = table.cell(0, j)
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.font.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            shading = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="{PECH_DARK_BLUE}" w:val="clear"/>'
            )
            cell._tc.get_or_add_tcPr().append(shading)
        return table

    # ---- BUILD DOCUMENT ----

    add_header()

    # SECTION 1
    add_section_heading("SECTION 1: PERSONAL INFORMATION")
    add_form_table([
        ("Surname / Last Name:", ""),
        ("First Name:", ""),
        ("Middle Name:", ""),
        ("Date of Birth:", "____/____/__________ (DD/MM/YYYY)"),
        ("Gender:", ["Male", "Female", "Prefer not to say"]),
        ("Marital Status:", ["Single", "Married", "Divorced", "Widowed"]),
        ("Nationality:", ""),
        ("State of Origin:", ""),
        ("LGA:", ""),
        ("Phone (Primary):", ""),
        ("Phone (Secondary):", ""),
        ("WhatsApp Number:", ""),
        ("Email Address:", ""),
        ("Residential Address:", ""),
        ("City / State:", ""),
        ("NIN:", ""),
    ])
    add_note("Please attach a recent passport-sized photograph.")

    # SECTION 2
    add_section_heading("SECTION 2: POSITION APPLIED FOR")
    add_form_table([
        ("Position / Job Title:", ""),
        ("Role ID (if known):", ""),
        ("Department:", ["Engineering", "Design", "Product & Business", "Operations", "Internship"]),
        ("Employment Type:", ["Full-Time", "Contract", "Internship"]),
        ("Expected Salary (NGN/month):", ""),
        ("Available Start Date:", "____/____/__________"),
        ("Notice Period:", ""),
        ("How did you hear about us?:", ""),
    ])

    # SECTION 3
    add_section_heading("SECTION 3: EDUCATIONAL BACKGROUND")
    add_note("List all qualifications starting from the most recent.")
    add_multi_row_table(["#", "Institution", "Qualification", "Course", "Year", "Grade"], 4)
    doc.add_paragraph()
    add_note("Professional Certifications:")
    add_multi_row_table(["#", "Certification", "Issuing Body", "Year", "Certificate No."], 3)

    # SECTION 4
    add_section_heading("SECTION 4: WORK EXPERIENCE")
    add_note("List your last 3 positions (most recent first). Write N/A if no experience.")
    for pos in range(1, 4):
        p = doc.add_paragraph()
        run = p.add_run(f"Position {pos}" + (" (Most Recent)" if pos == 1 else ""))
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x99, 0xCC)
        add_form_table([
            ("Company:", ""),
            ("Job Title:", ""),
            ("Type:", ["Full-Time", "Part-Time", "Contract", "Internship"]),
            ("Start Date:", "____/____/__________"),
            ("End Date:", "____/____/__________"),
            ("Still Employed?:", ["Yes", "No"]),
            ("Monthly Salary (NGN):", ""),
            ("Supervisor Name & Phone:", ""),
            ("Reason for Leaving:", ""),
            ("Key Responsibilities:", ""),
        ])
        doc.add_paragraph()

    # SECTION 5
    add_section_heading("SECTION 5: TECHNICAL SKILLS & COMPETENCIES")
    add_note("For technical roles. Non-technical applicants skip to General Competencies.")
    add_multi_row_table(["#", "Skill / Technology", "Proficiency (Beginner/Intermediate/Advanced/Expert)", "Years"], 8)
    doc.add_paragraph()
    add_form_table([
        ("Portfolio Website:", ""),
        ("GitHub Profile:", ""),
        ("LinkedIn Profile:", ""),
    ])

    # SECTION 6
    add_section_heading("SECTION 6: REFERENCES (Minimum 2)")
    for ref in range(1, 3):
        p = doc.add_paragraph()
        run = p.add_run(f"Reference {ref}")
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x99, 0xCC)
        add_form_table([
            ("Full Name:", ""),
            ("Relationship:", ""),
            ("Organisation & Position:", ""),
            ("Phone:", ""),
            ("Email:", ""),
        ])
        doc.add_paragraph()

    # SECTION 7
    add_section_heading("SECTION 7: GUARANTOR INFORMATION")
    add_note("Two guarantors required. Full Guarantor Forms completed upon offer.")
    for g in range(1, 3):
        p = doc.add_paragraph()
        run = p.add_run(f"Guarantor {g}")
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x99, 0xCC)
        add_form_table([
            ("Full Name:", ""),
            ("Occupation:", ""),
            ("Employer:", ""),
            ("Phone:", ""),
            ("Address:", ""),
        ])
        doc.add_paragraph()

    # SECTION 8
    add_section_heading("SECTION 8: HEALTH & EMERGENCY INFORMATION")
    add_form_table([
        ("Blood Group:", ["A+", "A-", "B+", "B-", "AB+", "AB-", "O+", "O-"]),
        ("Genotype:", ["AA", "AS", "SS", "AC", "SC"]),
        ("Medical Conditions?:", ["Yes", "No"]),
        ("If yes, specify:", ""),
        ("Disability/Accommodation?:", ["Yes", "No"]),
        ("If yes, specify:", ""),
        ("Currently on medication?:", ["Yes", "No"]),
    ])
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Emergency Contact")
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x99, 0xCC)
    add_form_table([
        ("Full Name:", ""),
        ("Relationship:", ""),
        ("Phone (Primary):", ""),
        ("Phone (Secondary):", ""),
        ("Address:", ""),
    ])

    # SECTION 9
    add_section_heading("SECTION 9: ADDITIONAL DECLARATIONS")
    add_form_table([
        ("Criminal conviction?:", ["Yes", "No"]),
        ("If yes, details:", ""),
        ("Previously dismissed?:", ["Yes", "No"]),
        ("If yes, details:", ""),
        ("Ongoing legal proceedings?:", ["Yes", "No"]),
        ("If yes, details:", ""),
        ("Relatives at PECH?:", ["Yes", "No"]),
        ("If yes, name & relationship:", ""),
        ("Currently employed?:", ["Yes", "No"]),
        ("If yes, employer & notice:", ""),
        ("Willing to relocate?:", ["Yes", "No", "Negotiable"]),
        ("Willing to travel?:", ["Yes", "No", "Domestic Only"]),
        ("Valid driver's license?:", ["Yes", "No"]),
        ("Valid passport?:", ["Yes", "No"]),
        ("Authorized to work in Nigeria?:", ["Yes", "No"]),
    ])

    # SECTION 10
    add_section_heading("SECTION 10: DECLARATION & CONSENT")
    p = doc.add_paragraph()
    run = p.add_run(
        "I declare that all information provided in this form is true, complete, and accurate. "
        "I consent to PECH Group Holdings Ltd collecting and processing my personal data in accordance "
        "with the Nigeria Data Protection Act (NDPA) 2023. I understand that any false information may "
        "result in disqualification or termination of employment."
    )
    run.font.size = Pt(9)
    doc.add_paragraph()

    # Data protection consent checkbox
    p = doc.add_paragraph()
    add_checkbox_sdt(p, "I consent to the processing of my personal data as described above")

    doc.add_paragraph()
    add_form_table([
        ("Applicant's Full Name:", ""),
        ("Signature:", ""),
        ("Date:", "____/____/__________ (DD/MM/YYYY)"),
    ])

    # SECTION 11
    add_section_heading("SECTION 11: FOR OFFICE USE ONLY")
    add_note("To be completed by HR Department only. Do NOT fill if you are the applicant.")
    add_form_table([
        ("Application Ref No.:", "CAF-________-________"),
        ("Date Received:", "____/____/__________"),
        ("Received By:", ""),
        ("Application Complete?:", ["Yes", "No"]),
        ("If No, missing docs:", ""),
    ])

    doc.add_paragraph()
    add_note("Document Verification Checklist:")
    verification_items = [
        "CV / Resume", "Cover Letter", "Passport Photograph",
        "Educational Certificates", "Professional Certifications",
        "NIN Verification", "Portfolio / GitHub (tech roles)"
    ]
    for item in verification_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"  {item}:  Received ")
        run.font.size = Pt(9)
        add_checkbox_sdt(p, "Yes")
        add_checkbox_sdt(p, "No")
        run2 = p.add_run("  Verified ")
        run2.font.size = Pt(9)
        add_checkbox_sdt(p, "Yes")
        add_checkbox_sdt(p, "No")
        add_checkbox_sdt(p, "N/A")

    doc.add_paragraph()
    add_note("Recommendation:")
    p = doc.add_paragraph()
    add_checkbox_sdt(p, "Shortlist for Interview — Proceed to Round 1")
    doc.add_paragraph()
    p = doc.add_paragraph()
    add_checkbox_sdt(p, "Waitlist — Qualified but hold for comparison")
    doc.add_paragraph()
    p = doc.add_paragraph()
    add_checkbox_sdt(p, "Reject — Does not meet minimum requirements")

    doc.add_paragraph()
    add_form_table([
        ("Reason for Decision:", ""),
        ("Interview Scheduled?:", ["Yes", "No"]),
        ("Round 1 Date:", "____/____/__________ at ______:______"),
        ("Round 1 Interviewer(s):", ""),
    ])

    doc.add_paragraph()
    add_note("Final Decision:")
    add_form_table([
        ("Final Status:", ["HIRED", "WAITLISTED", "REJECTED"]),
        ("Offer Made?:", ["Yes", "No"]),
        ("Offer Date:", "____/____/__________"),
        ("Offered Salary (NGN/month):", ""),
        ("Offer Accepted?:", ["Yes", "No"]),
        ("Start Date Confirmed:", "____/____/__________"),
    ])

    doc.add_paragraph()
    add_note("Approval Chain:")
    add_form_table([
        ("HR Officer:", "Name: ________________  Date: ____/____/____"),
        ("HR Officer Decision:", ["Approved", "Denied"]),
        ("Hiring Manager:", "Name: ________________  Date: ____/____/____"),
        ("Hiring Manager Decision:", ["Approved", "Denied"]),
        ("HR Head:", "Name: ________________  Date: ____/____/____"),
        ("HR Head Decision:", ["Approved", "Denied"]),
        ("MD/CEO (Senior Roles):", "Name: ________________  Date: ____/____/____"),
        ("MD/CEO Decision:", ["Approved", "Denied"]),
    ])

    # Footer
    add_gradient_divider()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(FOOTER_TAGLINE)
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    # Bottom gradient bar
    add_gradient_divider()

    # Save
    output_path = os.path.join(OUTPUT_DIR, "PECH_CANDIDATE_APPLICATION_FORM.docx")
    doc.save(output_path)
    print(f"Word file generated: {output_path}")
    return output_path


# ============================================================================
# MAIN
# ============================================================================

if __name__ == "__main__":
    print("=" * 60)
    print("PECH Group Holdings — Application Form Generator")
    print("=" * 60)
    print()

    excel_path = create_excel()
    print()
    word_path = create_word()

    print()
    print("=" * 60)
    print("Generation complete!")
    print(f"  Excel: {excel_path}")
    print(f"  Word:  {word_path}")
    print("=" * 60)
