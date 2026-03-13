#!/usr/bin/env python3
"""
PECH Group Holdings Ltd — XLSX to DOCX & PDF Converter

Converts Excel spreadsheets (.xlsx) to branded Word (.docx) and PDF files.
Reads all sheets from each xlsx, renders them as formatted tables in docx,
then generates a matching PDF.

Usage:
    python3 scripts/convert_xlsx_to_docx_pdf.py                    # Convert ALL xlsx files
    python3 scripts/convert_xlsx_to_docx_pdf.py path/to/file.xlsx  # Convert a single file

Brand Colors:
    Sky Blue:    #00BFFF  (primary)
    Orange:      #F5A623  (secondary/accent)
    Dark Blue:   #0099CC  (headings)
    Dark Orange: #E08A00  (H3 headings)
    Dark Navy:   #1B2838  (backgrounds)
"""

import os
import sys
import re
from pathlib import Path

import openpyxl
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

PROJECT_ROOT = Path(__file__).parent.parent

# ── Brand Colors ──────────────────────────────────────────────────────────────
SKY_BLUE = RGBColor(0x00, 0xBF, 0xFF)
DARK_BLUE = RGBColor(0x00, 0x99, 0xCC)
ORANGE = RGBColor(0xF5, 0xA6, 0x23)
DARK_ORANGE = RGBColor(0xE0, 0x8A, 0x00)
DARK_NAVY = RGBColor(0x1B, 0x28, 0x38)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BODY_TEXT = RGBColor(0x1A, 0x1A, 0x1A)

HEX_SKY_BLUE = "00BFFF"
HEX_DARK_BLUE = "0099CC"
HEX_ORANGE = "F5A623"
HEX_DARK_ORANGE = "E08A00"
HEX_DARK_NAVY = "1B2838"
HEX_WHITE = "FFFFFF"
HEX_LIGHT_BLUE = "E8F8FF"
HEX_BORDER_LIGHT = "B0D4E8"


def set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def create_branded_header(doc, title):
    """Add PECH branded header to document."""
    # Top ribbon
    body = doc.element.body

    ribbon_top = parse_xml(
        f'<w:p {nsdecls("w")}>'
        f'  <w:pPr>'
        f'    <w:pBdr>'
        f'      <w:top w:val="single" w:sz="36" w:space="0" w:color="{HEX_ORANGE}"/>'
        f'      <w:bottom w:val="single" w:sz="12" w:space="1" w:color="{HEX_SKY_BLUE}"/>'
        f'    </w:pBdr>'
        f'    <w:spacing w:after="0"/>'
        f'  </w:pPr>'
        f'</w:p>'
    )

    company = parse_xml(
        f'<w:p {nsdecls("w")}>'
        f'  <w:pPr>'
        f'    <w:jc w:val="center"/>'
        f'    <w:spacing w:before="200" w:after="40"/>'
        f'  </w:pPr>'
        f'  <w:r>'
        f'    <w:rPr>'
        f'      <w:b/>'
        f'      <w:color w:val="{HEX_DARK_BLUE}"/>'
        f'      <w:sz w:val="28"/>'
        f'      <w:szCs w:val="28"/>'
        f'    </w:rPr>'
        f'    <w:t>PECH GROUP HOLDINGS LTD</w:t>'
        f'  </w:r>'
        f'</w:p>'
    )

    tagline = parse_xml(
        f'<w:p {nsdecls("w")}>'
        f'  <w:pPr>'
        f'    <w:jc w:val="center"/>'
        f'    <w:spacing w:before="0" w:after="60"/>'
        f'  </w:pPr>'
        f'  <w:r>'
        f'    <w:rPr>'
        f'      <w:i/>'
        f'      <w:color w:val="{HEX_DARK_ORANGE}"/>'
        f'      <w:sz w:val="18"/>'
        f'      <w:szCs w:val="18"/>'
        f'    </w:rPr>'
        f'    <w:t>Technology &amp; Infrastructure Enablers for Africa</w:t>'
        f'  </w:r>'
        f'</w:p>'
    )

    thin_sep = parse_xml(
        f'<w:p {nsdecls("w")}>'
        f'  <w:pPr>'
        f'    <w:pBdr>'
        f'      <w:bottom w:val="single" w:sz="6" w:space="1" w:color="{HEX_SKY_BLUE}"/>'
        f'    </w:pBdr>'
        f'    <w:spacing w:before="0" w:after="120"/>'
        f'  </w:pPr>'
        f'</w:p>'
    )

    # Document title
    clean_title = title.replace("PECH_", "").replace("_", " ").title()
    title_para = parse_xml(
        f'<w:p {nsdecls("w")}>'
        f'  <w:pPr>'
        f'    <w:jc w:val="center"/>'
        f'    <w:spacing w:before="120" w:after="120"/>'
        f'  </w:pPr>'
        f'  <w:r>'
        f'    <w:rPr>'
        f'      <w:b/>'
        f'      <w:color w:val="{HEX_DARK_BLUE}"/>'
        f'      <w:sz w:val="32"/>'
        f'      <w:szCs w:val="32"/>'
        f'    </w:rPr>'
        f'    <w:t>{clean_title}</w:t>'
        f'  </w:r>'
        f'</w:p>'
    )

    ribbon_bottom = parse_xml(
        f'<w:p {nsdecls("w")}>'
        f'  <w:pPr>'
        f'    <w:pBdr>'
        f'      <w:top w:val="single" w:sz="12" w:space="1" w:color="{HEX_SKY_BLUE}"/>'
        f'      <w:bottom w:val="single" w:sz="36" w:space="0" w:color="{HEX_ORANGE}"/>'
        f'    </w:pBdr>'
        f'    <w:spacing w:before="0" w:after="240"/>'
        f'  </w:pPr>'
        f'</w:p>'
    )

    first_element = body[0] if len(body) > 0 else None
    elements = [ribbon_top, company, tagline, thin_sep, title_para, ribbon_bottom]
    if first_element is not None:
        for elem in reversed(elements):
            first_element.addprevious(elem)
    else:
        for elem in elements:
            body.append(elem)


def add_footer_ribbon(doc):
    """Add branded footer ribbon."""
    body = doc.element.body

    footer_sep = parse_xml(
        f'<w:p {nsdecls("w")}>'
        f'  <w:pPr>'
        f'    <w:pBdr>'
        f'      <w:top w:val="single" w:sz="36" w:space="1" w:color="{HEX_ORANGE}"/>'
        f'      <w:bottom w:val="single" w:sz="12" w:space="0" w:color="{HEX_SKY_BLUE}"/>'
        f'    </w:pBdr>'
        f'    <w:spacing w:before="300" w:after="0"/>'
        f'  </w:pPr>'
        f'</w:p>'
    )
    body.append(footer_sep)

    footer_text = parse_xml(
        f'<w:p {nsdecls("w")}>'
        f'  <w:pPr>'
        f'    <w:jc w:val="center"/>'
        f'    <w:spacing w:before="60" w:after="0"/>'
        f'  </w:pPr>'
        f'  <w:r>'
        f'    <w:rPr>'
        f'      <w:color w:val="{HEX_DARK_BLUE}"/>'
        f'      <w:sz w:val="16"/>'
        f'      <w:szCs w:val="16"/>'
        f'    </w:rPr>'
        f'    <w:t xml:space="preserve">PECH Group Holdings Ltd  |  Lagos, Nigeria  |  pechgroupholdings.tech</w:t>'
        f'  </w:r>'
        f'</w:p>'
    )
    body.append(footer_text)


def get_cell_text(cell):
    """Get text value from an openpyxl cell, handling None and numbers."""
    if cell.value is None:
        return ""
    return str(cell.value).strip()


def is_header_row(row, ws):
    """Heuristic: a row is a header if most cells have bold font or it's the first data row."""
    bold_count = 0
    filled_count = 0
    for cell in row:
        if cell.value is not None:
            filled_count += 1
            if cell.font and cell.font.bold:
                bold_count += 1
    return filled_count > 0 and bold_count >= filled_count * 0.5


def detect_data_region(ws):
    """
    Detect the main data table region in the worksheet.
    Returns (header_row_idx, data_start_row, data_end_row, start_col, end_col).
    """
    max_row = ws.max_row
    max_col = ws.max_column

    # Find the row with the most filled cells (likely header)
    best_row = 1
    best_count = 0
    for row_idx in range(1, min(max_row + 1, 30)):  # Check first 30 rows
        count = sum(1 for c in ws[row_idx] if c.value is not None)
        if count > best_count:
            best_count = count
            best_row = row_idx

    # Determine active columns from that row
    active_cols = []
    for cell in ws[best_row]:
        if cell.value is not None:
            active_cols.append(cell.column)

    if not active_cols:
        return None

    start_col = min(active_cols)
    end_col = max(active_cols)

    # Find data end (last row with any data in active columns)
    data_end = best_row
    for row_idx in range(best_row + 1, max_row + 1):
        has_data = False
        for col_idx in range(start_col, end_col + 1):
            if ws.cell(row=row_idx, column=col_idx).value is not None:
                has_data = True
                break
        if has_data:
            data_end = row_idx

    return (best_row, best_row + 1, data_end, start_col, end_col)


def extract_preamble_text(ws, header_row):
    """Extract any text before the main data table (titles, subtitles, etc.)."""
    preamble = []
    for row_idx in range(1, header_row):
        for cell in ws[row_idx]:
            if cell.value is not None:
                text = str(cell.value).strip()
                # Skip duplicate PECH branding (we add our own)
                if text in ("PECH GROUP HOLDINGS LTD",
                           "Technology & Infrastructure Enablers for People"):
                    continue
                if "pechgroupholdings.tech" in text.lower():
                    continue
                if text:
                    is_bold = cell.font and cell.font.bold
                    font_size = cell.font.size if cell.font and cell.font.size else 11
                    preamble.append((text, is_bold, font_size))
    return preamble


def extract_postamble_text(ws, data_end):
    """Extract any text after the main data table (notes, footers, etc.)."""
    postamble = []
    for row_idx in range(data_end + 1, ws.max_row + 1):
        for cell in ws[row_idx]:
            if cell.value is not None:
                text = str(cell.value).strip()
                if text in ("PECH GROUP HOLDINGS LTD",
                           "Technology & Infrastructure Enablers for People"):
                    continue
                if "pechgroupholdings.tech" in text.lower():
                    continue
                if "Confidential" in text and "PECH" in text:
                    continue
                if text:
                    postamble.append(text)
    return postamble


def xlsx_sheet_to_docx_table(doc, ws, sheet_name):
    """Convert an Excel worksheet to a formatted Word table."""
    region = detect_data_region(ws)
    if region is None:
        p = doc.add_paragraph(f"Sheet: {sheet_name} — (empty)")
        return

    header_row_idx, data_start, data_end, start_col, end_col = region
    num_cols = end_col - start_col + 1

    # Add preamble text
    preamble = extract_preamble_text(ws, header_row_idx)
    for text, is_bold, font_size in preamble:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.bold = is_bold
        if font_size and font_size > 14:
            run.font.size = Pt(14)
            run.font.color.rgb = DARK_BLUE
        elif is_bold:
            run.font.size = Pt(11)
            run.font.color.rgb = DARK_ORANGE

    # Create table
    num_rows = data_end - header_row_idx + 1
    if num_rows <= 0 or num_cols <= 0:
        return

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set table borders
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        tbl.insert(0, tblPr)

    borders_xml = (
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER_LIGHT}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER_LIGHT}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER_LIGHT}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER_LIGHT}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER_LIGHT}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER_LIGHT}"/>'
        f'</w:tblBorders>'
    )
    existing_borders = tblPr.find(qn('w:tblBorders'))
    if existing_borders is not None:
        tblPr.remove(existing_borders)
    tblPr.append(parse_xml(borders_xml))

    # Populate table
    for row_offset in range(num_rows):
        xlsx_row = header_row_idx + row_offset
        for col_offset in range(num_cols):
            xlsx_col = start_col + col_offset
            cell_value = ws.cell(row=xlsx_row, column=xlsx_col).value
            cell_text = str(cell_value) if cell_value is not None else ""

            docx_cell = table.cell(row_offset, col_offset)
            docx_cell.text = ""
            p = docx_cell.paragraphs[0]
            run = p.add_run(cell_text)

            if row_offset == 0:
                # Header row styling
                set_cell_shading(docx_cell, HEX_DARK_BLUE)
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(9)
            else:
                # Data row styling with alternating colors
                bg = HEX_LIGHT_BLUE if row_offset % 2 == 0 else HEX_WHITE
                set_cell_shading(docx_cell, bg)
                run.font.color.rgb = BODY_TEXT
                run.font.size = Pt(8.5)

    # Add postamble
    postamble = extract_postamble_text(ws, data_end)
    for text in postamble:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.color.rgb = BODY_TEXT


def convert_xlsx_to_docx(xlsx_path, docx_path):
    """Convert an Excel file to a branded Word document."""
    xlsx_path = Path(xlsx_path)
    docx_path = Path(docx_path)

    wb = openpyxl.load_workbook(str(xlsx_path), data_only=True)
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    title = xlsx_path.stem

    # Process each sheet
    for idx, sheet_name in enumerate(wb.sheetnames):
        ws = wb[sheet_name]

        if idx > 0:
            # Add page break between sheets
            doc.add_page_break()

        # Add sheet name as heading if multiple sheets
        if len(wb.sheetnames) > 1:
            heading = doc.add_heading(sheet_name, level=2)
            for run in heading.runs:
                run.font.color.rgb = DARK_BLUE
                run.font.bold = True

        xlsx_sheet_to_docx_table(doc, ws, sheet_name)

    # Add branded header and footer
    create_branded_header(doc, title)
    add_footer_ribbon(doc)

    doc.save(str(docx_path))
    return True


def convert_docx_to_pdf_via_html(docx_path, pdf_path):
    """Convert DOCX to PDF using pdfkit (wkhtmltopdf)."""
    import pdfkit

    # Read the docx and create HTML representation
    doc = Document(str(docx_path))

    html_parts = ["""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
    @page { margin: 1.5cm; size: A4; }
    body { font-family: Calibri, Arial, sans-serif; font-size: 10pt; color: #1A1A1A; margin: 0; padding: 20px; }
    .header-ribbon { border-top: 6px solid #F5A623; border-bottom: 2px solid #00BFFF; padding: 5px 0; margin-bottom: 10px; }
    .company-name { text-align: center; color: #0099CC; font-size: 14pt; font-weight: bold; margin: 10px 0 5px 0; }
    .tagline { text-align: center; color: #E08A00; font-size: 9pt; font-style: italic; margin: 0 0 10px 0; }
    .doc-title { text-align: center; color: #0099CC; font-size: 16pt; font-weight: bold; margin: 10px 0; }
    .separator { border-bottom: 1px solid #00BFFF; margin: 5px 0 15px 0; }
    .bottom-ribbon { border-top: 2px solid #00BFFF; border-bottom: 6px solid #F5A623; padding: 5px 0; margin-top: 20px; }
    .footer { text-align: center; color: #0099CC; font-size: 8pt; margin-top: 10px; }
    h2 { color: #0099CC; font-size: 13pt; border-bottom: 1px solid #00BFFF; padding-bottom: 3px; }
    table { width: 100%; border-collapse: collapse; margin: 10px 0; font-size: 8.5pt; }
    th { background-color: #0099CC; color: white; font-weight: bold; padding: 6px 4px; text-align: left; border: 1px solid #B0D4E8; font-size: 9pt; }
    td { padding: 5px 4px; border: 1px solid #B0D4E8; }
    tr:nth-child(even) td { background-color: #E8F8FF; }
    tr:nth-child(odd) td { background-color: #FFFFFF; }
    .preamble { margin: 5px 0; }
    .preamble-bold { font-weight: bold; color: #E08A00; }
    .preamble-title { font-size: 14pt; color: #0099CC; font-weight: bold; }
    .postamble { font-size: 9pt; margin: 5px 0; }
    .page-break { page-break-before: always; }
</style>
</head>
<body>
"""]

    # Extract title from document
    title = Path(docx_path).stem.replace("PECH_", "").replace("_", " ").title()

    html_parts.append(f"""
<div class="header-ribbon"></div>
<div class="company-name">PECH GROUP HOLDINGS LTD</div>
<div class="tagline">Technology &amp; Infrastructure Enablers for Africa</div>
<div class="separator"></div>
<div class="doc-title">{title}</div>
<div class="bottom-ribbon"></div>
""")

    # Process document content
    in_table = False
    for element in doc.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

        if tag == 'p':
            # Paragraph
            text_parts = []
            for child in element.iter():
                child_tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                if child_tag == 't':
                    text_parts.append(child.text or "")

            text = ''.join(text_parts).strip()
            if not text:
                continue

            # Check if it's a heading
            style_elem = element.find(qn('w:pPr'))
            if style_elem is not None:
                style_ref = style_elem.find(qn('w:pStyle'))
                if style_ref is not None:
                    style_val = style_ref.get(qn('w:val'), '')
                    if 'Heading' in style_val:
                        level = ''.join(c for c in style_val if c.isdigit()) or '2'
                        html_parts.append(f"<h{level}>{text}</h{level}>\n")
                        continue

            # Skip branded header/footer elements (already in HTML header)
            if text in ("PECH GROUP HOLDINGS LTD",
                       "Technology & Infrastructure Enablers for Africa"):
                continue
            if "pechgroupholdings.tech" in text.lower():
                continue

            html_parts.append(f"<p>{text}</p>\n")

        elif tag == 'tbl':
            # Table
            html_parts.append("<table>\n")
            rows = element.findall(qn('w:tr'))
            for row_idx, tr in enumerate(rows):
                html_parts.append("<tr>\n")
                cells = tr.findall(qn('w:tc'))
                cell_tag = "th" if row_idx == 0 else "td"
                for tc in cells:
                    cell_texts = []
                    for p in tc.findall(qn('w:p')):
                        for t in p.iter():
                            t_tag = t.tag.split('}')[-1] if '}' in t.tag else t.tag
                            if t_tag == 't':
                                cell_texts.append(t.text or "")
                    cell_text = ' '.join(cell_texts).strip()
                    html_parts.append(f"  <{cell_tag}>{cell_text}</{cell_tag}>\n")
                html_parts.append("</tr>\n")
            html_parts.append("</table>\n")

    # Footer
    html_parts.append("""
<div class="bottom-ribbon"></div>
<div class="footer">PECH Group Holdings Ltd  |  Lagos, Nigeria  |  pechgroupholdings.tech</div>
</body>
</html>
""")

    html_content = ''.join(html_parts)

    # Write temp HTML
    temp_html = str(pdf_path).replace('.pdf', '_temp.html')
    with open(temp_html, 'w', encoding='utf-8') as f:
        f.write(html_content)

    try:
        options = {
            'page-size': 'A4',
            'margin-top': '15mm',
            'margin-right': '15mm',
            'margin-bottom': '15mm',
            'margin-left': '15mm',
            'encoding': 'UTF-8',
            'no-outline': None,
            'enable-local-file-access': None,
            'quiet': '',
        }
        pdfkit.from_file(temp_html, str(pdf_path), options=options)
        return True
    except Exception as e:
        print(f"    PDF error: {e}")
        return False
    finally:
        try:
            os.unlink(temp_html)
        except OSError:
            pass


def convert_xlsx(xlsx_path):
    """Convert a single xlsx file to both docx and pdf."""
    xlsx_path = Path(xlsx_path)
    base = xlsx_path.stem
    out_dir = xlsx_path.parent

    docx_path = out_dir / f"{base}.docx"
    pdf_path = out_dir / f"{base}.pdf"

    print(f"  Converting: {xlsx_path.name}")

    try:
        # Generate DOCX
        convert_xlsx_to_docx(xlsx_path, docx_path)
        print(f"    + DOCX: {docx_path.name}")

        # Generate PDF from DOCX
        if convert_docx_to_pdf_via_html(docx_path, pdf_path):
            print(f"    + PDF:  {pdf_path.name}")
        else:
            print(f"    x PDF:  Failed")

        return True
    except Exception as e:
        print(f"    x FAILED: {e}")
        return False


def find_all_xlsx():
    """Find all xlsx files that need conversion."""
    xlsx_files = []

    # Root level
    for f in PROJECT_ROOT.glob("*.xlsx"):
        xlsx_files.append(f)

    # employment_xlsx directory
    emp_dir = PROJECT_ROOT / "employment_xlsx"
    if emp_dir.exists():
        for f in emp_dir.glob("*.xlsx"):
            xlsx_files.append(f)

    return sorted(xlsx_files)


def main():
    print("=" * 65)
    print("  PECH XLSX → DOCX + PDF Converter")
    print("  Generating branded Word and PDF from Excel spreadsheets")
    print("=" * 65)
    print()

    if len(sys.argv) > 1:
        files = [Path(f) for f in sys.argv[1:]]
    else:
        files = find_all_xlsx()

    if not files:
        print("No XLSX files found to convert.")
        return

    print(f"Found {len(files)} XLSX files to convert.\n")

    success = 0
    failed = 0
    for f in files:
        if convert_xlsx(f):
            success += 1
        else:
            failed += 1

    print()
    print("=" * 65)
    print(f"  Done! {success} converted, {failed} failed")
    print("=" * 65)


if __name__ == "__main__":
    main()
